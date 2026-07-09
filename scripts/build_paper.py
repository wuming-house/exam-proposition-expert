#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_paper.py — 命题专家技能配套脚本（自包含，无需 pandoc）

把按约定写好的 Markdown 试卷（含 LaTeX 公式与图片引用）转换为
可直接打印的 Word 试卷（.docx）。

原理：
  1. 轻量 Markdown 解析器把试卷拆成标题/段落/列表/表格/图片/公式块。
  2. LaTeX 公式经 `latex2mathml` 解析后，渲染为 Unicode 数学符号（分数、根号、上/下标、
     希腊字母等）插入 Word 正文，因此公式可在 Word / WPS / LibreOffice 中清晰显示，
     不是图片，也不会乱码。
  3. 图片用 python-docx 直接嵌入。
  4. 后处理统一中文字体、A4 页面、页边距、页码，可选左侧"密封线"。

依赖：python-docx、latex2mathml（可选 matplotlib 用于生成配图）。
用法：
  python build_paper.py paper.md -o 试卷.docx [--seamless] [--no-page-number]
"""

import argparse
import os
import re
import sys
from lxml import etree as lxml_et
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Mm, Emu
from docx.text.paragraph import Paragraph

try:
    from latex2mathml.converter import convert as latex_to_mathml
except Exception as e:  # pragma: no cover
    sys.exit(f"[错误] 未安装 latex2mathml：请先 `pip install latex2mathml`。\n({e})")

# 注册命名空间，确保序列化时使用标准前缀（m:/wp:/a:/wps:），
# 否则 lxml 会生成 ns0 等前缀，虽然 Word 仍能识别，但统一前缀更稳妥。
try:
    from lxml import etree as _lxml_etree
    for _p, _u in {
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    }.items():
        _lxml_etree.register_namespace(_p, _u)
except Exception:
    pass

# 中文字体方案（标准中文试卷排版）
BODY_ASCII = "Times New Roman"
BODY_EA = "宋体"
HEAD_EA = "黑体"
TITLE_SIZE = Pt(16)
SECTION_SIZE = Pt(13)
SUBSECTION_SIZE = Pt(11.5)
BODY_SIZE = Pt(10.5)


# ----------------------------------------------------------------------------
# 工具：命名空间
# ----------------------------------------------------------------------------
def m(local):
    """创建 OMML(math) 命名空间元素。"""
    return OxmlElement(f"m:{local}")


def localname(tag):
    if not isinstance(tag, str):
        return None
    return tag.split("}", 1)[1] if "}" in tag else tag


# ----------------------------------------------------------------------------
def _fix_degree_symbol(omath):
    """后处理：把上标/重音中的 ∘(U+2218 RING OPERATOR) 替换为 °(U+00B0 DEGREE SIGN)。

    LaTeX 的 \\circ 在数学排版中产生 U+2218（环运算符），
    但「角度/度数」场景下应该用 U+00B0（度符号）。
    U+2218 在非 Word 环境（IMA/WPS/LibreOffice 等）中常被渲染为
    大圆圈或附带多余字符，而 U+00B0 是通用标准字符，跨平台一致。
    """
    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    RING = "\u2218"   # ∘  RING OPERATOR
    DEGREE = "\u00b0" # °  DEGREE SIGN

    for t in omath.iter(f"{{{M_NS}}}t"):
        if t.text and RING in t.text:
            t.text = t.text.replace(RING, DEGREE)
    return omath


# MathML 结构规整化
# ----------------------------------------------------------------------------
_CLOSING = {")": "(", "]": "[", "}": "{", "\u27e9": "\u27e8",
            "\u2309": "\u2308", "\u230b": "\u230a"}  # 右→左 括号映射

_MATH = "http://www.w3.org/1998/Math/MathML"


def _normalize_mathml(math_elem):
    """规整 latex2mathml 输出的畸形 MathML 结构。

    已知问题：latex2mathml 对 ``(-2)^3`` 生成
      <mrow><mo>(</mo><mo>−</mo><mn>2</mn>
        <msup><mo>)</mo><mn>3</mn></msup></mrow>
    即右括号被塞进 msup 底数、其余内容散落在 mrow 中。
    正确结构应为整个 (-2) 作为 msup 底数。

    本函数扫描 msup/msub/msubsup，若底数为右括号类字符则向前回溯
    匹配的左括号，将中间所有元素包裹进 <mrow> 作为新底数。
    """
    import copy

    for tag in ("msup", "msub", "msubsup"):
        for node in math_elem.iter(f"{{{_MATH}}}{tag}"):
            children = list(node)
            if not children:
                continue
            base = children[0]  # msup/msub: base; msubsup: base
            base_text = (base.text or "").strip()
            # 检查底数是否是单个右括号类字符
            if base_text not in _CLOSING and len(list(base)) == 0:
                # 也检查只有文本子节点的 mo 元素
                t_nodes = list(base.iter(f"{{{_MATH}}}t"))
                if not t_nodes or (t_nodes[0].text or "").strip() not in _CLOSING:
                    continue
                base_text = (t_nodes[0].text or "").strip()

            opening = _CLOSING.get(base_text)
            if opening is None:
                continue

            parent = node.getparent()
            if parent is None:
                continue
            siblings = list(parent)

            # 找 node 在 siblings 中的位置和匹配的左括号
            idx = siblings.index(node)
            match_idx = None
            depth = 0
            # 从 node 前面向后找（跳过嵌套的同类型标签）
            for i in range(idx - 1, -1, -1):
                sib = siblings[i]
                sib_tag = localname(sib.tag)
                sib_txt = "".join(sib.itertext()).strip()
                if sib_tag == tag:
                    depth += 1
                elif depth > 0 and (sib_tag == "mrow" and any(localname(c.tag) == tag for c in sib)):
                    depth -= 1
                elif depth == 0 and sib_txt == opening:
                    match_idx = i
                    break

            if match_idx is None:
                continue

            # 把 siblings[match_idx .. idx-1] + 原底数包进 <mrow> 当新底数
            new_base = lxml_et.Element(f"{{{_MATH}}}mrow")
            for i in range(match_idx, idx):
                new_base.append(copy.deepcopy(siblings[i]))
            new_base.append(copy.deepcopy(base))  # deepcopy！lxml.append会移动元素

            # 替换底数
            node.remove(base)
            node.insert(0, new_base)

            # 删除原位置上被吸进去的兄弟节点（从后往前删避免索引漂移）
            for i in range(idx - 1, match_idx - 1, -1):
                parent.remove(siblings[i])

    return math_elem


# LaTeX -> Word 原生方程（OMML）
# ----------------------------------------------------------------------------
# amsmath 环境兼容表：latex2mathml 不支持 aligned / eqnarray（会泄漏 & 对齐符，
# 生成非法 MathML 导致 XML 解析失败），但它支持 align。把前者改写为 align 即可
# 正确生成多行矩阵（m:m），保留 & 列对齐语义。align/gather/cases/matrix 等原生支持。
_ALIGN_REPL = [
    (r"\begin{aligned}", r"\begin{align}"),
    (r"\end{aligned}", r"\end{align}"),
    (r"\begin{aligned*}", r"\begin{align*}"),
    (r"\end{aligned*}", r"\end{align*}"),
    (r"\begin{eqnarray}", r"\begin{align}"),
    (r"\end{eqnarray}", r"\end{align}"),
    (r"\begin{eqnarray*}", r"\begin{align*}"),
    (r"\end{eqnarray*}", r"\end{align*}"),
]


def _preprocess_latex(tex):
    """把 latex2mathml 不支持的对齐环境改写为其支持的等价环境。"""
    for a, b in _ALIGN_REPL:
        tex = tex.replace(a, b)
    return tex


def latex_to_omath(latex):
    """把一段 LaTeX 转成 <m:oMath> 元素；失败则退化为纯文本。"""
    try:
        mathml = latex_to_mathml(_preprocess_latex(latex))
        # 用 lxml 解析（支持 .getparent()，规整化需要）
        root = lxml_et.fromstring(mathml.encode("utf-8"))
        _normalize_mathml(root)   # 修复 latex2mathml 的畸形结构
        omath = mathml_to_omath(root)
        return _fix_degree_symbol(omath)  # 度数符号标准化
    except Exception as e:
        sys.stderr.write(f"[提示] 公式转换失败，已退化为文本：{latex} ({e})\n")
        o = m("oMath")
        r = m("r")
        rpr_fallback = m("rPr")
        rf_fb = m("rFonts")
        rf_fb.set(qn("m:ascii"), "Cambria Math")
        rf_fb.set(qn("m:hAnsi"), "Cambria Math")
        rpr_fallback.append(rf_fb)
        r.append(rpr_fallback)
        t = m("t")
        t.text = latex
        r.append(t)
        o.append(r)
        return o


# -*- coding: utf-8 -*-
"""
Correct v5 _latex_to_unicode function.
This file will be inserted into build_paper.py between the start/end markers.
"""

# ====== 全面 Unicode 降级（v5：覆盖分数/根号/对齐环境，避免 OMML 在 WPS 中变方框）======

_SUB_MAP = {
    '0': '\u2080', '1': '\u2081', '2': '\u2082', '3': '\u2083',
    '4': '\u2084', '5': '\u2085', '6': '\u2086', '7': '\u2087',
    '8': '\u2088', '9': '\u2089', '+': '\u208a', '-': '\u208b',
    '=': '\u208c', '(': '\u208d', ')': '\u208e',
}
_SUP_MAP = {
    '0': '\u2070', '1': '\u00b9', '2': '\u00b2', '3': '\u00b3',
    '4': '\u2074', '5': '\u2075', '6': '\u2076', '7': '\u2077',
    '8': '\u2078', '9': '\u2079', '+': '\u207a', '-': '\u207b',
    '=': '\u207c', '(': '\u207d', ')': '\u207e', 'n': '\u207f',
}

_LATEX_SYMBOL = {}

# Populate symbol table safely
_sym_data = [
    ("\u2192", "\u2192"),  # placeholder
]
# We'll populate this below using direct dict assignment

import re

# Build _LATEX_SYMBOL by direct assignment (avoid escape issues)
_pairs = [
    # Each tuple: (laTeX_cmd_string, unicode_char)
    # Using chr() concatenation for backslash + cmd name
]
# Actually, let me just use a different approach - define commands as variables

# LaTeX command names (each is backslash + letters)  
# Use explicit string construction to avoid escaping hell
_bslash = chr(92)  # backslash character

_CMD_FRAC   = _bslash + "frac"
_CMD_DFRAC  = _bslash + "dfrac" 
_CMD_TFRAC  = _bslash + "tfrac"
_CMD_BEGIN  = _bslash + "begin"
_CMD_END    = _bslash + "end"
_CMD_SQRT   = _bslash + "sqrt"
_CMD_LEFT   = _bslash + "left"
_CMD_RIGHT  = _bslash + "right"
_CMD_BIGL   = _bslash + "bigl"
_CMD_BIGR   = _bslash + "bigr"
_CMD_BIGL2  = _bslash + "Bigl"
_CMD_BIGR2  = _bslash + "Bigr"
_CMD_BIGGL  = _bslash + "biggl"
_CMD_BIGGR  = _bslash + "biggr"
_CMD_TEXT   = _bslash + "text"
_CMD_MBF    = _bslash + "mathbf"
_CMD_MRM    = _bslash + "mathrm"
_CMD_MBB    = _bslash + "mathbb"
_CMD_MCAL   = _bslash + "mathcal"
_CMD_XRIGHT = _bslash + "xrightarrow"
_CMD_XLEFT  = _bslash + "xleftarrow"
_CMD_QUAD   = _bslash + "quad"
_CMD_QQUAD  = _bslash + "qquad"

_LATEX_SYMBOL = {
    _bslash + "rightarrow": "\u2192", _bslash + "to": "\u2192",
    _bslash + "leftarrow": "\u2190", _bslash + "Leftarrow": "\u21d0",
    _bslash + "Rightarrow": "\u21d2", _bslash + "Leftrightarrow": "\u21d4",
    _bslash + "leftrightarrow": "\u2194", _bslash + "leq": "\u2264",
    _bslash + "rightleftharpoons": "\u21cc", _bslash + "leftrightharpoons": "\u21cc",
    _bslash + "geq": "\u2265", _bslash + "neq": "\u2260", _bslash + "pm": "\u00b1",
    _bslash + "approx": "\u2248", _bslash + "times": "\u00d7", _bslash + "cdot": "\u22c5",
    _bslash + "div": "\u00f7", _bslash + "ldots": "\u2026", _bslash + "cdots": "\u22ef",
    _bslash + "uparrow": "\u2191", _bslash + "downarrow": "\u2193",
    _bslash + "alpha": "\u03b1", _bslash + "beta": "\u03b2", _bslash + "gamma": "\u03b3",
    _bslash + "delta": "\u03b4", _bslash + "theta": "\u03b8", _bslash + "pi": "\u03c0",
    _bslash + "sigma": "\u03c3", _bslash + "omega": "\u03c9", _bslash + "Delta": "\u0394",
    _bslash + "Sigma": "\u03a3", _bslash + "Omega": "\u03a9", _bslash + "infty": "\u221e",
    _bslash + "circ": "\u00b0", _bslash + "degree": "\u00b0",
    _bslash + "uplus": "\u228e", _bslash + "cap": "\u2229", _bslash + "cup": "\u222a",
    _bslash + "in": "\u2208", _bslash + "notin": "\u2209", _bslash + "subset": "\u2282",
    _bslash + "supset": "\u2283", _bslash + "subseteq": "\u2286", _bslash + "supseteq": "\u2287",
    _bslash + "emptyset": "\u2205", _bslash + "nabla": "\u2207", _bslash + "partial": "\u2202",
    _bslash + "prime": "\u2032", _bslash + "sqrt": "\u221a",
    # Large operators (approximate as Unicode base symbol)
    _bslash + "sum": "\u2211",       # ∑
    _bslash + "int": "\u222b",       # ∫
    _bslash + "prod": "\u220f",      # ∏
    _bslash + "lim": "lim",
    _bslash + "over": "/",
    _bslash + "choose": "C",
    _bslash + "stackrel": "\u2192",  # simplified arrow
    _bslash + "overset": "^",
    _bslash + "underset": "_",
    
        # Accents/combining symbols
    _bslash + "bar": "\u0304",    # combining macron/overline
    _bslash + "hat": "\u0302",    # combining circumflex
    _bslash + "vec": "\u20d7",    # combining right arrow above
    _bslash + "dot": "\u0307",    # combining dot above
    _bslash + "tilde": "\u0303",  # combining tilde
    _bslash + "overline": "\u203e",
}

_OMML_ONLY_PATTERNS = [
    r'\\begin\{(matrix|cases|bmatrix|pmatrix|vmatrix|Vmatrix)\}',
    r'\\sum_\{', r'\\int_\{', r'\\prod_\{', r'\\lim_\{',
    r'\\over\s*\{', r'\\choose\s*\{',
    r'\\stackrel\{', r'\\overset\{', r'\\underset\{',
]


def _parse_matrix_table(latex):
    """Parse \\begin{matrix|array|...}...\\end{...} into list-of-lists (cells already Unicode-converted).
    Returns None if latex is not a matrix environment."""
    m = re.match(r'\\begin\{(matrix|array|bmatrix|pmatrix|vmatrix|Vmatrix)\}\s*(.*?)\\end\{\1\}', latex, re.DOTALL)
    if not m:
        return None
    body = m.group(2).strip()
    rows = []
    for line in body.split(_bslash * 2):
        line = line.strip()
        if not line:
            continue
        cells = [c.strip() for c in line.split('&')]
        uni_cells = []
        for c in cells:
            u = _latex_to_unicode(c)
            uni_cells.append(u if u is not None else c)
        rows.append(uni_cells)
    return rows



def _is_omml_only(latex):
    """v6: Always return False - all LaTeX is converted to Unicode or Word tables.
    No formula should ever use OMML (causes □ boxes in user's software)."""
    return False


def _find_matching_brace(latex, start):
    """Find matching } from position start (which should be at {)."""
    depth = 0
    i = start
    while i < len(latex):
        if latex[i] == '{':
            depth += 1
        elif latex[i] == '}':
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return len(latex) - 1


def _latex_to_unicode(latex):
    """Convert LaTeX formula to plain Unicode text (v5 comprehensive).
    
    Supports subscripts, superscripts, fractions (\\frac, \\dfrac, \\tfrac),
    square roots (\\sqrt), aligned environments (\\begin{aligned}...\\end{aligned}),
    common symbols (\\times, \\cdot, \\div, \\rightarrow, etc.), \\text{}, \\mathbf{},
    \\left(...\\right), and more.
    
    Returns None only for truly unconvertible structures like matrices,
    summation/integral limits, etc.
    """
    if _is_omml_only(latex):
        return None
    
    result = []
    i = 0
    n = len(latex)
    
    while i < n:
        ch = latex[i]
        
        if ch == _bslash and i + 1 < n:
            
            # ====== \frac{num}{den} / \dfrac / \tfrac ======
            frac_matched = False
            for fcmd, flen in [(_CMD_FRAC, len(_CMD_FRAC)), 
                                (_CMD_DFRAC, len(_CMD_DFRAC)),
                                (_CMD_TFRAC, len(_CMD_TFRAC))]:
                if latex[i:i+flen] == fcmd:
                    j = i + flen
                    while j < n and latex[j] in ' \t':
                        j += 1
                    if j < n and latex[j] == '{':
                        num_end = _find_matching_brace(latex, j)
                        num_inner = latex[j+1:num_end]
                        k = num_end + 1
                        while k < n and latex[k] in ' \t':
                            k += 1
                        if k < n and latex[k] == '{':
                            den_end = _find_matching_brace(latex, k)
                            den_inner = latex[k+1:den_end]
                            num_u = _latex_to_unicode(num_inner) or num_inner
                            den_u = _latex_to_unicode(den_inner) or den_inner
                            result.append(num_u + '/' + den_u)
                            i = den_end + 1
                            frac_matched = True
                            break
                    if not frac_matched:
                        result.append('frac')
                        i = flen  # wrong: should advance past cmd
                        # Fix: we need to advance i past the command
                        # But we've already consumed up to i+flen conceptually
                        # Actually we haven't moved i yet since we're still at the backslash position
                        i += flen
                        frac_matched = True
                    break
            
            if frac_matched:
                continue
            
            # ====== \begin{aligned} ... \end{aligned} ======
            if latex[i:i+len(_CMD_BEGIN)] == _CMD_BEGIN and i+len(_CMD_BEGIN) < n and latex[i+len(_CMD_BEGIN)] == '{':
                bstart = i + len(_CMD_BEGIN) + 1
                bend = latex.index('}', bstart)
                env_name = latex[bstart:bend].strip()
                aligned_names = ['aligned', 'align', 'aligned*', 'align*']
                if env_name in aligned_names:
                    base_name = env_name.rstrip('*')
                    end_found = False
                    for try_name in [env_name, base_name]:
                        end_tag = _CMD_END + '{' + try_name + '}'
                        end_pos = latex.find(end_tag, bend + 1)
                        if end_pos != -1:
                            body = latex[bend+1:end_pos]
                            body_lines = body.split(_bslash + _bslash)  # split on \\
                            converted_lines = []
                            for bl in body_lines:
                                bl = bl.strip()
                                if bl.startswith('&'):
                                    bl = bl[1:].strip()
                                bl = bl.replace('&', '  ')
                                bl = re.sub(r'\\(?:quad|qquad)', '  ', bl)
                                cline = _latex_to_unicode(bl) or bl
                                converted_lines.append(cline)
                            result.append('\n'.join(converted_lines))
                            i = end_pos + len(end_tag)
                            end_found = True
                            break
                    if end_found:
                        continue
                else:
                    # matrix/cases → not convertible
                    return None
            
            # ====== \end{...} isolated ======
            if latex[i:i+len(_CMD_END)] == _CMD_END and i+len(_CMD_END) < n and latex[i+len(_CMD_END)] == '{':
                eend = latex.index('}', i + len(_CMD_END) + 1)
                i = eend + 1
                continue
            
            # ====== \xrightarrow{...} / \xleftarrow{...} ======
            xr_matched = False
            for xrcmd, xrarrow, xrlen in [
                (_CMD_XRIGHT, '\u2192', len(_CMD_XRIGHT)),
                (_CMD_XLEFT, '\u2190', len(_CMD_XLEFT)),
            ]:
                if latex[i:i+xrlen] == xrcmd:
                    start = i + xrlen
                    if start < n and latex[start] == '{':
                        end = _find_matching_brace(latex, start)
                        i = end + 1
                    else:
                        i += xrlen
                    result.append(xrarrow)
                    xr_matched = True
                    break
            if xr_matched:
                continue
            
            # ====== \sqrt{...} ======
            if latex[i:i+len(_CMD_SQRT)] == _CMD_SQRT:
                j = i + len(_CMD_SQRT)
                if j < n and latex[j] == '{':
                    end = _find_matching_brace(latex, j)
                    inner = latex[j+1:end]
                    inner_u = _latex_to_unicode(inner) or inner
                    result.append('\u221a' + inner_u)
                    i = end + 1
                elif j < n and latex[j] == '[':
                    bracket_end = latex.index(']', j + 1)
                    k = bracket_end + 1
                    while k < n and latex[k] in ' \t':
                        k += 1
                    if k < n and latex[k] == '{':
                        end = _find_matching_brace(latex, k)
                        inner = latex[k+1:end]
                        inner_u = _latex_to_unicode(inner) or inner
                        result.append('\u221a' + inner_u)
                        i = end + 1
                    else:
                        result.append('\u221a')
                        i = k
                else:
                    result.append('\u221a')
                    i = j
                continue
            
            # ====== \left / \right / \bigl / \bigr etc ======
            delim_cmds = [_CMD_LEFT, _CMD_RIGHT, _CMD_BIGL, _CMD_BIGR, 
                         _CMD_BIGL2, _CMD_BIGR2, _CMD_BIGGL, _CMD_BIGGR]
            delim_matched = False
            for dc in delim_cmds:
                dclen = len(dc)
                if latex[i:i+dclen] == dc:
                    j = i + dclen
                    # Only treat as delimiter command if followed by actual delimiter
                    if j < n and (latex[j] in '([{.)]}|' or latex[j] == _bslash or latex[j] == '.'):
                        if latex[j] in '([{.)]}|':
                            result.append(latex[j])
                        elif latex[j] == _bslash and j+1 < n:
                            nc = latex[j+1]
                            dm = {'{': '}', '}': '{', '.': ''}
                            result.append(dm.get(nc, nc))
                            j += 1
                        elif latex[j] == '.':
                            pass  # empty delimiter
                        i = j + 1
                        delim_matched = True
                    break
            if delim_matched:
                continue
            
            # ====== \text{...}, \mathbf{...}, \mathrm{...}, \mathbb{...}, \mathcal{...} ======
            font_cmds = [(_CMD_TEXT, len(_CMD_TEXT)), (_CMD_MBF, len(_CMD_MBF)),
                         (_CMD_MRM, len(_CMD_MRM)), (_CMD_MBB, len(_CMD_MBB)),
                         (_CMD_MCAL, len(_CMD_MCAL))]
            font_matched = False
            for fc, fclen in font_cmds:
                if latex[i:i+fclen] == fc and i+fclen < n and latex[i+fclen] == '{':
                    j = latex.index('}', i + fclen + 1)
                    result.append(latex[i+fclen+1:j])
                    i = j + 1
                    font_matched = True
                    break
            if font_matched:
                continue
            
            # ====== \\ (newline) ======
            if latex[i+1] == _bslash:
                result.append('\n')
                i += 2
                continue
            
            # ====== \ (space) ======
            if latex[i+1] == ' ':
                result.append(' ')
                i += 2
                continue
            
            # ====== \quad / \qquad ======
            if latex[i:i+len(_CMD_QUAD)] == _CMD_QUAD:
                result.append('  ')
                i += len(_CMD_QUAD)
                continue
            if latex[i:i+len(_CMD_QQUAD)] == _CMD_QQUAD:
                result.append('    ')
                i += len(_CMD_QQUAD)
                continue
            
            # ====== spacing: \, \: \; \! ======
            if latex[i+1] in (',', ':', ';', '!'):
                result.append(' ')
                i += 2
                continue
            
            # ====== Known symbols from dict (longest match first) ======
            sym_matched = False
            for sym_cmd in sorted(_LATEX_SYMBOL.keys(), key=len, reverse=True):
                if latex[i:].startswith(sym_cmd):
                    end = i + len(sym_cmd)
                    if end < n and latex[end].isalpha():
                        continue
                    result.append(_LATEX_SYMBOL[sym_cmd])
                    i = end
                    sym_matched = True
                    break
            if sym_matched:
                continue
            
            # ====== Fallback: unrecognized command ======
            result.append(latex[i + 1] if i + 1 < n else '')
            i += 2
        
        elif ch == '_' and i + 1 < n:
            if latex[i + 1] == '{':
                j = _find_matching_brace(latex, i + 1)
                inner = latex[i + 2:j]
                # Recurse: subscript content may contain \text{}, commands, etc.
                inner_u = _latex_to_unicode(inner) or inner
                # Convert result to subscript characters where possible
                for c in inner_u:
                    result.append(_SUB_MAP.get(c, c))
                i = j + 1
            else:
                result.append(_SUB_MAP.get(latex[i + 1], latex[i + 1]))
                i += 2
        
        elif ch == '^' and i + 1 < n:
            if latex[i + 1] == '{':
                j = _find_matching_brace(latex, i + 1)
                inner = latex[i + 2:j]
                # Recurse: superscript content may contain complex LaTeX
                inner_u = _latex_to_unicode(inner) or inner
                for c in inner_u:
                    result.append(_SUP_MAP.get(c, c))
                i = j + 1
            else:
                result.append(_SUP_MAP.get(latex[i + 1], latex[i + 1]))
                i += 2
        
        else:
            result.append(ch)
            i += 1
    
    return ''.join(result)


_is_simple_latex = lambda latex: not _is_omml_only(latex)
_simple_to_unicode = _latex_to_unicode



def _children_e(parent_omml, mm_parent):
    """把 MathML 子节点转换后追加进 OMML 父节点（mrow/mstyle 直接展开）。"""
    for child in list(mm_parent):
        ln = localname(child.tag)
        if ln is None:
            continue
        if ln in ("mrow", "mstyle"):
            _children_e(parent_omml, child)
        elif ln in ("mphantom",):
            # 占位符：忽略内容
            continue
        else:
            parent_omml.append(_convert(child))


def mathml_to_omath(math_root):
    omath = m("oMath")
    # math_root 可能是 <math> 或已是表达式
    _children_e(omath, math_root)
    return omath


def _make_run(elem):
    ln = localname(elem.tag)
    text = "".join(elem.itertext())
    r = m("r")
    sty_val = None
    mv = elem.get("mathvariant")
    if mv:
        sty_val = {"italic": "i", "bold": "b", "bold-italic": "bi",
                   "normal": "p", "double-struck": "p", "script": "p"}.get(mv)
    else:
        if ln == "mi":
            sty_val = "i"
        elif ln == "mtext":
            sty_val = "p"
    # 每个 OMML run 必须显式声明 m:rFonts=Cambria Math，
    # 否则 Word/WPS 会从 Normal 样式继承 Times New Roman 等
    # 非数学字体 → 上下标/根号等布局异常 → 显示为方框。
    rpr = m("rPr")
    rf = m("rFonts")
    rf.set(qn("m:ascii"), "Cambria Math")
    rf.set(qn("m:hAnsi"), "Cambria Math")
    rpr.append(rf)
    if sty_val:
        sty = m("sty")
        sty.set(qn("m:val"), sty_val)
        rpr.append(sty)
    r.append(rpr)
    t = m("t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _e_arg(mm_elem):
    """生成 m:e 参数节点（用于分组 / 作为 f、sup 等的参数）。"""
    e = m("e")
    _children_e(e, mm_elem)
    return e


def _unwrap_e(elem):
    """若 elem 是 <m:e>，返回其子元素列表（拆包）；否则返回 [elem]。

    用于 msup/msub 等的 base 参数：_convert(mrow) 会返回 <m:e>，
    而外层 handler 又创建了 base=<m:e>，导致双层嵌套。
    此函数消除多余的一层。
    """
    M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    if elem is not None and localname(elem.tag) == "e":
        return list(elem)
    return [elem]


def _e_from(child):
    """把【单个】MathML 元素转成 <m:e> 包裹的 OMML。

    正确处理叶子（mn/mi/mo）与容器（mrow/mfrac/msup 等）：
      _children_e 期望传入「容器」并迭代其 children，若传入叶子（无 children）
      则什么都不填 —— 这正是之前上标/分数/根号内容丢失的根因。
    """
    e = m("e")
    e.append(_convert(child))
    return e


def _convert(elem):
    ln = localname(elem.tag)
    kids = [c for c in list(elem) if localname(c.tag) is not None]

    if ln in ("mi", "mn", "mo", "mtext"):
        return _make_run(elem)
    if ln == "mspace":
        r = m("r")
        rpr = m("rPr")
        rf = m("rFonts")
        rf.set(qn("m:ascii"), "Cambria Math")
        rf.set(qn("m:hAnsi"), "Cambria Math")
        rpr.append(rf)
        r.append(rpr)
        t = m("t")
        t.set(qn("xml:space"), "preserve")
        t.text = " "
        r.append(t)
        return r
    if ln in ("mrow", "mstyle"):
        return _e_arg(elem)
    if ln == "mfrac":
        f = m("f")
        f.append(m("fPr"))
        num = m("num")
        den = m("den")
        if len(kids) >= 1:
            num.append(_e_from(kids[0]))
        if len(kids) >= 2:
            den.append(_e_from(kids[1]))
        f.append(num)
        f.append(den)
        return f
    if ln == "msqrt":
        rad = m("rad")
        rad.append(m("radPr"))
        deg = m("deg")
        rad.append(deg)
        e = m("e")
        if kids:
            e.append(_convert(kids[0]))
        rad.append(e)
        return rad
    if ln == "mroot":
        rad = m("rad")
        rad.append(m("radPr"))
        deg = m("deg")
        e = m("e")
        if len(kids) >= 2:
            deg.append(_convert(kids[1]))
            e.append(_convert(kids[0]))
        rad.append(deg)
        rad.append(e)
        return rad
    if ln == "msup":
        wrap = m("sSup")       # 上标函数容器（不是 m:sup！）
        base = m("e")
        exp = m("sup")         # 上标内容
        if len(kids) >= 1:
            for child in _unwrap_e(_convert(kids[0])):
                base.append(child)
        if len(kids) >= 2:
            exp.append(_convert(kids[1]))
        wrap.append(base)
        wrap.append(exp)
        return wrap
    if ln == "msub":
        wrap = m("sSub")       # 下标函数容器（不是 m:sub！）
        base = m("e")
        sub = m("sub")         # 下标内容
        if len(kids) >= 1:
            for child in _unwrap_e(_convert(kids[0])):
                base.append(child)
        if len(kids) >= 2:
            sub.append(_convert(kids[1]))
        wrap.append(base)
        wrap.append(sub)
        return wrap
    if ln == "msubsup":
        wrap = m("subSup")
        base = m("e")
        sub = m("sub")
        sup = m("sup")
        if len(kids) >= 1:
            for child in _unwrap_e(_convert(kids[0])):
                base.append(child)
        if len(kids) >= 2:
            sub.append(_convert(kids[1]))
        if len(kids) >= 3:
            sup.append(_convert(kids[2]))
        wrap.append(base)
        wrap.append(sub)
        wrap.append(sup)
        return wrap
    if ln == "mfenced":
        d = m("d")
        dpr = m("dPr")
        beg = m("begChr")
        beg.set(qn("m:val"), elem.get("open", "(") or "(")
        end = m("endChr")
        end.set(qn("m:val"), elem.get("close", ")") or ")")
        dpr.append(beg)
        dpr.append(end)
        d.append(dpr)
        e = m("e")
        for c in kids:
            e.append(_convert(c))
        d.append(e)
        return d
    if ln == "mtable":
        tbl = m("m")
        tbl.append(m("mPr"))
        for row in kids:
            if localname(row.tag) != "mtr":
                continue
            mr = m("mr")
            cells = [c for c in list(row) if localname(c.tag) == "mtd"]
            for cell in cells:
                e = m("e")
                _children_e(e, cell)
                mr.append(e)
            tbl.append(mr)
        return tbl
    if ln in ("mover", "munder"):
        acc = m("acc")
        accpr = m("accPr")
        # 取 mo 作为重音符号
        mo_text = ""
        for c in kids:
            if localname(c.tag) == "mo":
                mo_text = "".join(c.itertext())
                break
        chr_el = m("chr")
        chr_el.set(qn("m:val"), mo_text or "^")
        accpr.append(chr_el)
        acc.append(accpr)
        base = m("e")
        if kids:
            base.append(_convert(kids[0]))
        acc.append(base)
        return acc
    if ln == "mphantom":
        # 占位：返回空 e
        return m("e")
    # 兜底：当作分组
    return _e_arg(elem)


# ----------------------------------------------------------------------------
# 中文字体 / 排版
# ----------------------------------------------------------------------------
def set_run_fonts(run, ascii_font, ea_font, size=None):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), ea_font)
    rfonts.set(qn("w:cs"), ascii_font)


def set_style_font(style, ascii_font, ea_font, size=None, bold=None, align=None):
    try:
        rpr = style.element.get_or_add_rPr()
    except Exception:
        return
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), ea_font)
    rfonts.set(qn("w:cs"), ascii_font)
    if size is not None:
        style.font.size = size
    if bold is not None:
        style.font.bold = bold
    if align is not None:
        style.paragraph_format.alignment = align


def style_document(doc):
    for style in doc.styles:
        try:
            set_style_font(style, BODY_ASCII, BODY_EA)
        except Exception:
            pass
    names = [s.name for s in doc.styles]
    if "Normal" in names:
        set_style_font(doc.styles["Normal"], BODY_ASCII, BODY_EA, size=BODY_SIZE)
    if "Heading 1" in names:
        set_style_font(doc.styles["Heading 1"], BODY_ASCII, HEAD_EA,
                       size=TITLE_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if "Heading 2" in names:
        set_style_font(doc.styles["Heading 2"], BODY_ASCII, HEAD_EA,
                       size=SECTION_SIZE, bold=True)
    if "Heading 3" in names:
        set_style_font(doc.styles["Heading 3"], BODY_ASCII, HEAD_EA,
                       size=SUBSECTION_SIZE, bold=True)
    # 正文 run 补中文
    for p in doc.paragraphs:
        for r in p.runs:
            if r._element.find(qn("m:oMath")) is not None:
                continue
            set_run_fonts(r, BODY_ASCII, BODY_EA)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_run_fonts(r, BODY_ASCII, BODY_EA)


def set_page_setup(doc):
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_field(paragraph, field):
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field} "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(b)
    r._r.append(instr)
    r._r.append(sep)
    r._r.append(end)


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p.add_run("第 ")
    set_run_fonts(r0, BODY_ASCII, BODY_EA, Pt(9))
    add_field(p, "PAGE")
    r1 = p.add_run(" 页 / 共 ")
    set_run_fonts(r1, BODY_ASCII, BODY_EA, Pt(9))
    add_field(p, "NUMPAGES")
    r2 = p.add_run(" 页")
    set_run_fonts(r2, BODY_ASCII, BODY_EA, Pt(9))


def add_seam_line(doc):
    try:
        from docx.oxml.ns import nsmap
        nsmap.setdefault("wps",
            "http://schemas.microsoft.com/office/word/2010/wordprocessingShape")
        nsmap.setdefault("wpg",
            "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup")
        nsmap.setdefault("wpc",
            "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingCanvas")

        section = doc.sections[0]
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        pos_left = Emu(int(Cm(0.4)))
        pos_top = Emu(int(Cm(105)))
        box_w = Emu(int(Cm(0.8)))
        box_h = Emu(int(Cm(85)))

        anchor = OxmlElement("wp:anchor")
        for k, v in {
            "distT": "0", "distB": "0", "distL": "0", "distR": "0",
            "simplePos": "0", "relativeHeight": "251658240",
            "behindDoc": "1", "locked": "0", "layoutInCell": "1",
            "allowOverlap": "1",
        }.items():
            anchor.set(qn("wp:" + k), v)
        sp = OxmlElement("wp:simplePos")
        sp.set(qn("wp:x"), "0")
        sp.set(qn("wp:y"), "0")
        anchor.append(sp)
        ph = OxmlElement("wp:positionH")
        ph.set(qn("wp:relativeFrom"), "page")
        ph_off = OxmlElement("wp:posOffset")
        ph_off.text = str(int(pos_left))
        ph.append(ph_off)
        anchor.append(ph)
        pv = OxmlElement("wp:positionV")
        pv.set(qn("wp:relativeFrom"), "page")
        pv_off = OxmlElement("wp:posOffset")
        pv_off.text = str(int(pos_top))
        pv.append(pv_off)
        anchor.append(pv)
        ext = OxmlElement("wp:extent")
        ext.set(qn("wp:cx"), str(int(box_w)))
        ext.set(qn("wp:cy"), str(int(box_h)))
        anchor.append(ext)
        eff = OxmlElement("wp:effectExtent")
        for k, v in {"l": "0", "t": "0", "r": "0", "b": "0"}.items():
            eff.set(qn("wp:" + k), v)
        anchor.append(eff)
        anchor.append(OxmlElement("wp:wrapNone"))
        docpr = OxmlElement("wp:docPr")
        docpr.set(qn("wp:id"), "99")
        docpr.set(qn("wp:name"), "SeamLine")
        anchor.append(docpr)
        anchor.append(OxmlElement("wp:cNvGraphicFramePr"))

        graphic = OxmlElement("a:graphic")
        gdata = OxmlElement("a:graphicData")
        gdata.set(qn("a:uri"),
                  "http://schemas.microsoft.com/office/word/2010/wordprocessingShape")
        wsp = OxmlElement("wps:wsp")
        txbx = OxmlElement("wps:txbx")
        content = OxmlElement("w:txbxContent")
        para = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        td = OxmlElement("w:textDirection")
        td.set(qn("w:val"), "tbRl")
        ppr.append(td)
        para.append(ppr)
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:eastAsia"), "宋体")
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        rpr.append(sz)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "808080")
        rpr.append(color)
        r.append(rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = "密 封 线"
        r.append(t)
        para.append(r)
        content.append(para)
        txbx.append(content)
        wsp.append(txbx)
        gdata.append(wsp)
        graphic.append(gdata)
        anchor.append(graphic)

        drawing = OxmlElement("w:drawing")
        drawing.append(anchor)
        run = hp.add_run()
        run._r.append(drawing)
    except Exception as e:
        sys.stderr.write(f"[提示] 密封线插入失败，已跳过（不影响试卷主体）：{e}\n")


# ----------------------------------------------------------------------------
# 轻量 Markdown 解析 -> docx
# ----------------------------------------------------------------------------
# 行内分词：公式($...$)优先级最高，其次「**$公式$**」加粗公式，
# 再次普通粗体(**...**)。粗体内若含公式则递归解析，确保公式不被吞成文本。
INLINE_RE = re.compile(r"\*\*\$([^$]+)\$\*\*|\$([^$]+)\$|\*\*([^*]+)\*\*")
BLANK_RE = re.compile(r"_{2,}")


def _fill_cell_with_blanks(paragraph, text):
    """把文本中的连续下划线（如 ________）渲染为等长空格的 Word 下划线。

    避免某些字体/缩放比例下，下划线字符显示过短或被截断的观感问题。
    """
    pos = 0
    for mm in BLANK_RE.finditer(text):
        if mm.start() > pos:
            r = paragraph.add_run(text[pos:mm.start()])
            set_run_fonts(r, BODY_ASCII, BODY_EA)
        # 用等长空格的 Word 下划线替代下划线字符
        blank_len = mm.end() - mm.start()
        r = paragraph.add_run(" " * blank_len)
        r.underline = True
        set_run_fonts(r, BODY_ASCII, BODY_EA)
        pos = mm.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_fonts(r, BODY_ASCII, BODY_EA)


def tokenize_inline(s):
    tokens = []
    pos = 0
    for mm in INLINE_RE.finditer(s):
        if mm.start() > pos:
            tokens.append(("text", s[pos:mm.start()]))
        if mm.group(1) is not None:
            # **$公式$** —— 加粗公式（公式优先渲染）
            tokens.append(("math_bold", mm.group(1)))
        elif mm.group(2) is not None:
            tokens.append(("math", mm.group(2)))
        else:
            # 普通粗体；若其内部含公式则递归解析，避免公式退化为文本
            inner = mm.group(3)
            if "$" in inner:
                tokens.extend(tokenize_inline(inner))
            else:
                tokens.append(("bold", inner))
        pos = mm.end()
    if pos < len(s):
        tokens.append(("text", s[pos:]))
    return tokens


def fill_inline(paragraph, text):
    has_math = False
    for kind, val in tokenize_inline(text):
        if kind == "text":
            if val == "":
                continue
            r = paragraph.add_run(val)
            set_run_fonts(r, BODY_ASCII, BODY_EA)
        elif kind == "bold":
            r = paragraph.add_run(val)
            r.bold = True
            set_run_fonts(r, BODY_ASCII, BODY_EA)
        else:  # math / math_bold —— 行内公式（`**$...$**` 也按公式渲染）
            has_math = True
            unicode_text = _simple_to_unicode(val)
            if unicode_text is not None:
                # 简单公式（仅下标/上标）：用 Unicode 纯文本，
                # 避免 OMML 在 WPS/旧版 Word 中渲染为方框
                r = paragraph.add_run(unicode_text)
                if kind == "math_bold":
                    r.bold = True
                set_run_fonts(r, BODY_ASCII, BODY_EA)
            else:
                # v6 fallback: raw text (no OMML - causes boxes)
                r = paragraph.add_run(val)
                if kind == "math_bold":
                    r.bold = True
                set_run_fonts(r, BODY_ASCII, BODY_EA)
    # 包含内联公式的段落：强制设置行距 + 段前/段后间距，
    # 防止上标/度数符号（超出基线以上的部分）被行高顶部裁掉。
    # 原理：Word 的 atLeast 行距通常把富余空间加在基线下方，
    # 而上标/度数符号的溢出在基线*上方*，故额外加 spaceBefore 给顶部留白。
    if has_math:
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        pf.line_spacing = Pt(22)
        pf.space_before = Pt(3)
        pf.space_after = Pt(2)


# 参考答案专用配色：答案=蓝、解析/知识点=绿
COLOR_ANSWER = "1F4E79"    # 深蓝
COLOR_ANALYSIS = "2E7D32"  # 深绿


def _color_paragraph(paragraph, hex_color):
    """给段落内所有文字上色（含公式 m:r 内的文字）。"""
    try:
        from docx.shared import RGBColor
        for r in paragraph.runs:
            r.font.color.rgb = RGBColor.from_string(hex_color)
    except Exception:
        pass
    # 公式内 <m:r> 上色（OMML 用 m:color）
    try:
        for mr in paragraph._p.iter(qn("m:r")):
            rpr = mr.find(qn("m:rPr"))
            if rpr is None:
                rpr = OxmlElement("m:rPr")
                mr.insert(0, rpr)
            col = rpr.find(qn("m:color"))
            if col is None:
                col = OxmlElement("m:color")
                rpr.append(col)
            col.set(qn("m:val"), hex_color)
    except Exception:
        pass


def fill_colored_line(doc, text, hex_color):
    """添加一行文本并整体着色（用于参考答案的答案/解析行）。"""
    p = doc.add_paragraph()
    fill_inline(p, text)
    _color_paragraph(p, hex_color)
    return p


def _insert_colored_after(paragraph, text, hex_color):
    """在指定段落后插入一个彩色文本段落（用于答案/解析嵌入题目下方）。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    fill_inline(new_para, text)
    _color_paragraph(new_para, hex_color)
    return new_para


def parse_answers(answer_text):
    """解析参考答案文本，返回 {题号: (答案文本, 解析文本)}。

    支持格式：
      1. **C** — 解析内容
      2. **答案内容** — 解析内容
      11. **$CaO$** — 氧化钙。
      19. **（12 分）** $n = 0.2$ mol。
      多题一行：10. **B**  11. **C**  12. **B**
      表格答案：| 题号 | 1 | 2 | ... | 答案 | C | C | ...
    """
    answers = {}
    if not answer_text or not answer_text.strip():
        return answers

    lines = answer_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # 跳过 Markdown 标题行
        if re.match(r"^#{1,3}\s+", line):
            i += 1
            continue

        # 处理表格：表头 | 题号 | 1 | 2 | ... ; 数据行 | 答案 | C | C | ...
        if line.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # 过滤分隔行
            data_rows = [row for row in table_lines if not re.match(r"^\|\s*[:\-]+\s*\|", row)]
            if len(data_rows) >= 2:
                header = [c.strip() for c in data_rows[0].strip("|").split("|")]
                for row in data_rows[1:]:
                    cells = [c.strip() for c in row.strip("|").split("|")]
                    if not cells:
                        continue
                    first = cells[0]
                    if first in ("答案", "答案解析", "参考答案"):
                        for idx, val in enumerate(cells[1:], start=1):
                            if idx < len(header):
                                try:
                                    qnum = int(header[idx])
                                    answers[qnum] = (val, "")
                                except ValueError:
                                    pass
            continue

        # 处理一行多题：10. **B**  11. **C**  12. **B**
        # 只有当一行出现 ≥2 个 "数字. **答案**" 对时才按多题处理，
        # 避免把 1. **C** — 解析 这种单题行误判。
        pairs = list(re.finditer(r"(\d+)\.\s*\*\*(.+?)\*\*", line))
        if len(pairs) >= 2:
            for m in pairs:
                qnum = int(m.group(1))
                ans = m.group(2).strip()
                answers[qnum] = (ans, "")
            i += 1
            continue

        # 处理单题：1. **C** — 解析
        m = re.match(r"^(\d+)\.\s*(.*)$", line)
        if not m:
            i += 1
            continue

        qnum = int(m.group(1))
        rest = m.group(2).strip()

        ans_match = re.match(r"^\*\*(.+?)\*\*\s*(.*)$", rest)
        if ans_match:
            ans = ans_match.group(1).strip()
            analysis = ans_match.group(2).strip()
        else:
            # 无 **答案** 包裹：若整行是评分要点/解析/思路/知识点，则作为解析
            if re.match(r"^(评分要点|解析|思路|知识点|注意)", rest):
                ans = ""
                analysis = rest
            else:
                ans = rest
                analysis = ""

        # 去掉解析前导的分隔符（—、--、：等）
        analysis = re.sub(r"^[\u2014\u2013\u2012\-\:\：]\s*", "", analysis)

        # 合并后续非空行（只要没有新题号/表格/标题/多题模式）。
        # 用换行 \n 连接，保留多段结构，渲染时可拆成多个绿色段落，
        # 像老师板书一样分层讲解（【思路】【步骤】【知识点】【易错】等）。
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                continue
            if (re.match(r"^\d+\.\s*", nxt) or nxt.startswith("|")
                    or re.match(r"^#{1,3}\s+", nxt)
                    or len(list(re.finditer(r"\d+\.\s*\*\*", nxt))) >= 2):
                break
            analysis += ("\n" + nxt) if analysis else nxt
            i += 1

        answers[qnum] = (ans, analysis)

    return answers


def is_table_start(lines, i):
    if "|" not in lines[i]:
        return False
    nxt = lines[i + 1].strip() if i + 1 < len(lines) else ""
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", nxt)) and "-" in nxt


def remove_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tblPr.append(borders)


def _set_equal_column_widths(tbl, cols):
    """设置表格所有列为等宽（基于页面可用宽度）。"""
    try:
        # 获取/创建 tblGrid
        tbl_grid = tbl._tbl.find(qn("w:tblGrid"))
        if tbl_grid is None:
            tbl_grid = OxmlElement("w:tblGrid")
            tbl._tbl.insert(0, tbl_grid)
        else:
            tbl_grid.clear()
        # 页面可用宽度 = 页宽 - 左边距 - 右边距
        page_width_emu = int(Mm(210))  # A4 页宽
        left_margin_emu = int(Cm(2.5))
        right_margin_emu = int(Cm(2.5))
        available = page_width_emu - left_margin_emu - right_margin_emu
        col_width = available // max(cols, 1)
        for _ in range(cols):
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(col_width))
            tbl_grid.append(gc)
    except Exception as e:
        sys.stderr.write(f"[提示] 表格等宽设置失败，已跳过：{e}\n")


def build_docx(md_text, md_dir, mode='paper', answers=None, seamless=False, no_page_number=False):
    """把 Markdown 试卷文本转为 docx。

    mode:
      'paper'  -> 学生卷（仅题目，不含答案）
      'answer' -> 教师卷（题目 + 彩色答案 + 彩色解析）
    """
    doc = Document()
    lines = md_text.split("\n")
    n = len(lines)
    i = 0

    # 在答案模式下记录每道题题干段落，以便后面插入答案/解析
    qnum_to_para = {}

    while i < n:
        line = lines[i]
        strip = line.strip()
        if strip == "":
            i += 1
            continue
        if strip.startswith("#"):
            level = len(strip) - len(strip.lstrip("#"))
            title = strip[level:].strip()
            doc.add_heading(title, level=min(level, 3))
            i += 1
        elif strip in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
        elif re.match(r"^!\[[^\]]*\]\([^)]+\)$", strip):
            m = re.match(r"^!\[[^\]]*\]\(([^)]+)\)$", strip)
            path = m.group(1).strip()
            if not os.path.isabs(path):
                path = os.path.join(md_dir, path)
            if os.path.isfile(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p.add_run().add_picture(path, width=Cm(9))
                except Exception:
                    p.add_run(f"[图片缺失: {path}]")
            else:
                p = doc.add_paragraph()
                p.add_run(f"[图片缺失: {path}]")
                set_run_fonts(p.runs[0], BODY_ASCII, BODY_EA)
            i += 1
        elif strip.startswith("$$"):
            # 展示公式：支持跨行 $$...$$（如 \begin{align} 多行推导）。
            # 注意：单独一行的 $$ 是多行块的开头，需向后收集直到闭合 $$。
            if strip == "$$":
                # 开头是单独的 $$，开始跨行收集（首行无内容）
                buf = []
                i += 1
                closed = False
                while i < n:
                    cur = lines[i].strip()
                    if cur.endswith("$$"):
                        if cur != "$$":
                            buf.append(cur[:-2])
                        closed = True
                        i += 1
                        break
                    buf.append(lines[i])
                    i += 1
                latex = "\n".join(buf).strip()
            elif strip.endswith("$$"):
                # 单行块：同一行内 $...$$
                latex = strip[2:-2].strip()
                i += 1
            else:
                # 跨行块：开头行带有部分内容（如 $$ \begin{align}）
                buf = [strip[2:]]
                i += 1
                closed = False
                while i < n:
                    cur = lines[i].strip()
                    if cur.endswith("$$"):
                        buf.append(cur[:-2])
                        closed = True
                        i += 1
                        break
                    buf.append(lines[i])
                    i += 1
                latex = "\n".join(buf).strip()
                if not closed:
                    # 未找到闭合 $$：退化为普通段落，避免吞掉内容
                    p = doc.add_paragraph()
                    p.add_run("\n".join(buf))
                    set_run_fonts(p.runs[0], BODY_ASCII, BODY_EA)
                    continue
            if latex == "":
                # 空公式（无内容）：跳过，不生成空段落
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # v6: Try matrix first (-> Word native table)
            matrix_data = _parse_matrix_table(latex)
            if matrix_data is not None:
                n_rows = len(matrix_data)
                n_cols = max((len(row) for row in matrix_data), default=1)
                tbl = doc.add_table(rows=n_rows, cols=n_cols)
                try:
                    tbl.style = "Table Grid"
                except Exception:
                    pass
                tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for ri, row_data in enumerate(matrix_data):
                    for ci, cell_text in enumerate(row_data):
                        cell = tbl.cell(ri, ci)
                        cell.text = ""
                        p_inner = cell.paragraphs[0]
                        p_inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p_inner.add_run(cell_text)
                        set_run_fonts(r, BODY_ASCII, BODY_EA)
            else:
                unicode_text = _simple_to_unicode(latex)
                if unicode_text is not None:
                    r = p.add_run(unicode_text)
                    set_run_fonts(r, BODY_ASCII, BODY_EA)
                else:
                    # Ultimate fallback: raw text (should never happen in v6)
                    r = p.add_run(latex)
                    set_run_fonts(r, BODY_ASCII, BODY_EA)
            i += 1
        elif is_table_start(lines, i):
            # 解析表格
            header = [c.strip() for c in strip.strip().strip("|").split("|")]
            i += 2  # 跳过分隔行
            body = []
            while i < n and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                body.append(row)
                i += 1
            cols = max(len(header), *(len(r) for r in body), 1)
            tbl = doc.add_table(rows=1, cols=cols)
            try:
                tbl.style = "Table Grid"
            except Exception:
                pass
            # 表头
            for c in range(cols):
                cell = tbl.rows[0].cells[c]
                cell.text = ""
                _fill_cell_with_blanks(cell.paragraphs[0], header[c] if c < len(header) else "")
            for r in body:
                cells = tbl.add_row().cells
                for c in range(cols):
                    cells[c].text = ""
                    _fill_cell_with_blanks(cells[c].paragraphs[0], r[c] if c < len(r) else "")
            # 考生信息栏（含下划线填空）去边框 + 等宽列
            joined = " ".join(header + [x for r in body for x in r])
            if "____" in joined:
                remove_table_borders(tbl)
                # 强制等宽：每列宽度 = 页面可用宽度 / 列数
                _set_equal_column_widths(tbl, cols)
        elif (strip.startswith("【答案】") or re.match(r"^答案[：:]", strip)
              or strip.startswith("**答案**")):
            # 参考答案：答案行（蓝色），写在题目下方
            fill_colored_line(doc, line, COLOR_ANSWER)
            i += 1
        elif (strip.startswith("【解析】") or strip.startswith("【知识点】")
              or strip.startswith("【思路】")
              or re.match(r"^(解析|知识点|思路)[：:]", strip)):
            # 参考答案：解析/知识点行（绿色），写在答案下方
            fill_colored_line(doc, line, COLOR_ANALYSIS)
            i += 1
        elif re.match(r"^\d+\.\s", strip) or strip.startswith("- ") or strip.startswith("* "):
            # 列表（试卷题号通常以 "1. " 形式出现）
            items = []
            while i < n and lines[i].strip():
                s = lines[i].strip()
                mm = re.match(r"^(\d+)\.\s+(.*)$", s)
                if mm:
                    items.append(("num", int(mm.group(1)), mm.group(2)))
                    i += 1
                elif s.startswith("- ") or s.startswith("* "):
                    items.append(("bul", None, s[2:]))
                    i += 1
                else:
                    break
            for kind, qnum, content in items:
                style = "List Number" if kind == "num" else "List Bullet"
                try:
                    p = doc.add_paragraph(style=style)
                except Exception:
                    p = doc.add_paragraph()
                fill_inline(p, content)
                if kind == "num" and mode == "answer":
                    qnum_to_para[qnum] = p
        else:
            p = doc.add_paragraph()
            fill_inline(p, line)
            i += 1

    if mode == "answer" and answers:
        # 把答案/解析插入到每道题的末尾（下一道题之前），而不是题干后。
        # 用段落底层 XML 元素 _p 做映射，避免 Paragraph 包装对象不一致。
        p_to_qnum = {p._p: q for q, p in qnum_to_para.items()}
        qnum_idx = []
        for idx, para in enumerate(doc.paragraphs):
            if para._p in p_to_qnum:
                qnum_idx.append((p_to_qnum[para._p], idx))

        insertions = []
        for i, (qnum, start_idx) in enumerate(qnum_idx):
            if i + 1 < len(qnum_idx):
                end_idx = qnum_idx[i + 1][1] - 1
            else:
                end_idx = len(doc.paragraphs) - 1
            if qnum in answers:
                insertions.append((qnum, end_idx))

        # 从后往前插入，避免前面插入影响后续段落位置。
        # 由于 _insert_colored_after 用 addnext 在当前段落后插入，
        # 先插解析、再插答案，最终顺序才是：题目 -> 答案 -> 解析。
        for qnum, end_idx in sorted(insertions, key=lambda x: x[1], reverse=True):
            para = doc.paragraphs[end_idx]
            ans, analysis = answers[qnum]
            # 解析按行拆成多个绿色段落，像老师板书分层讲解。
            # 已有【思路】【步骤】【知识点】【易错】等自有标签的行不再加前缀；
            # 其余行自动加【解析】前缀，作为讲解的主干。
            if analysis:
                for sub in reversed([s for s in analysis.split("\n") if s.strip()]):
                    sub = sub.strip()
                    prefix = "" if sub.startswith("【") else "【解析】"
                    _insert_colored_after(para, f"{prefix}{sub}", COLOR_ANALYSIS)
            if ans:
                _insert_colored_after(para, f"【答案】{ans}", COLOR_ANSWER)

    style_document(doc)
    set_page_setup(doc)
    if not no_page_number:
        add_page_number_footer(doc)
    if seamless:
        add_seam_line(doc)
    return doc


def main():
    ap = argparse.ArgumentParser(description="Markdown 试卷 -> 可打印 docx（自包含）")
    ap.add_argument("input", help="输入的 paper.md 路径")
    ap.add_argument("-o", "--output", required=True, help="输出 docx 路径（自动生成 _试卷 和 _答案 两个文档）")
    ap.add_argument("--seamless", action="store_true", help="添加左侧竖排'密封线'")
    ap.add_argument("--no-page-number", action="store_true", help="不加页码")
    args = ap.parse_args()

    if not os.path.isfile(args.input):
        sys.exit(f"[错误] 找不到输入文件：{args.input}")

    with open(args.input, "r", encoding="utf-8") as f:
        md_text = f.read()

    md_dir = os.path.dirname(os.path.abspath(args.input))

    # 分割题目部分与参考答案部分
    split_idx = len(md_text)
    for m in re.finditer(r"^#{1,3}\s*参考答案", md_text, re.MULTILINE):
        split_idx = m.start()
        break

    question_text = md_text[:split_idx]
    answer_text = md_text[split_idx:]

    # 解析参考答案
    answers = parse_answers(answer_text)
    if answers:
        print(f"[0/3] 检测到参考答案，共 {len(answers)} 道题...")
    else:
        print("[0/3] 未检测到标准参考答案，答案卷将只含题目...")

    # 生成两个输出路径
    out_path = os.path.abspath(args.output)
    out_dir = os.path.dirname(out_path)
    base_name, ext = os.path.splitext(os.path.basename(out_path))
    paper_path = os.path.join(out_dir, f"{base_name}_试卷{ext}")
    answer_path = os.path.join(out_dir, f"{base_name}_答案{ext}")

    os.makedirs(out_dir, exist_ok=True)

    print("[1/3] 生成学生卷（仅题目）...")
    paper_doc = build_docx(question_text, md_dir, mode="paper",
                           seamless=args.seamless, no_page_number=args.no_page_number)

    print("[2/3] 生成教师卷（题目 + 彩色答案 + 彩色解析）...")
    answer_doc = build_docx(question_text, md_dir, mode="answer", answers=answers,
                              seamless=args.seamless, no_page_number=args.no_page_number)

    print("[3/3] 保存文档...")
    paper_doc.save(paper_path)
    answer_doc.save(answer_path)
    print(f"  -> 学生卷：{paper_path}")
    print(f"  -> 教师卷：{answer_path}")


if __name__ == "__main__":
    main()
